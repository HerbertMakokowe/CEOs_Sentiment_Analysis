"""
Generate Word Document Report for CEO Sentiment Analysis
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def create_report():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('CEO Sentiment Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Sentiment Analysis of US Top Company CEO Communications')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents a comprehensive sentiment analysis of CEO shareholder letters '
        'and communications from the top 10 US companies by market capitalization. Using advanced '
        'natural language processing techniques, we analyzed how these business leaders communicate '
        'about building their companies, the challenges they faced, and their strategic vision.'
    )
    
    # What We Did Section
    doc.add_heading('Project Overview', level=1)
    doc.add_paragraph(
        'We built an end-to-end sentiment analysis pipeline that performs the following steps:'
    )
    
    steps = [
        ('Data Collection', 'Scraped CEO shareholder letters from SEC EDGAR 10-K filings for 10 major US companies including Apple, Microsoft, Amazon, Nvidia, Alphabet, Meta, Tesla, Berkshire Hathaway, JPMorgan, and ExxonMobil.'),
        ('Text Preprocessing', 'Cleaned and processed the raw text by removing boilerplate legal language, forward-looking statement disclaimers, and other non-relevant content. Chunked text into meaningful paragraph-level segments.'),
        ('Keyword Filtering', 'Filtered paragraphs containing keywords related to company building: built, journey, founded, growth, challenge, vision, strategy, milestone, decision, learned.'),
        ('Triple Sentiment Analysis', 'Applied three different sentiment models to each text chunk for comprehensive analysis.'),
        ('Visualization', 'Generated 5 professional visualizations and a detailed markdown report summarizing the findings.'),
    ]
    
    for title_text, desc in steps:
        p = doc.add_paragraph()
        p.add_run(f'• {title_text}: ').bold = True
        p.add_run(desc)
    
    # Sentiment Models Section
    doc.add_heading('Sentiment Analysis Models', level=1)
    
    models = [
        ('VADER', 'A rule-based sentiment analyzer optimized for social media text. Provides a compound score from -1 (negative) to +1 (positive). Fast and efficient for large datasets.'),
        ('TextBlob', 'A pattern-based sentiment analyzer that provides polarity (-1 to +1) and subjectivity (0 to 1) scores. Good for general-purpose text analysis.'),
        ('FinBERT', 'A state-of-the-art transformer model (BERT) fine-tuned specifically on financial text. Provides sentiment classification (Positive/Neutral/Negative) with confidence scores. Most accurate for SEC filings and financial communications.'),
    ]
    
    for model_name, model_desc in models:
        p = doc.add_paragraph()
        p.add_run(f'{model_name}: ').bold = True
        p.add_run(model_desc)
    
    # Why It's Important Section
    doc.add_heading('Why This Analysis Matters', level=1)
    
    importance_points = [
        ('Investor Insights', 'CEO sentiment in shareholder letters can be a leading indicator of company performance. Positive, confident language often correlates with strong financial results, while cautious or negative sentiment may signal challenges ahead.'),
        ('Market Sentiment', 'Understanding how top CEOs communicate helps investors gauge overall market sentiment and business confidence among industry leaders.'),
        ('Communication Patterns', 'Analyzing multiple companies reveals communication patterns and benchmarks. Investors can compare how different leaders frame challenges and opportunities.'),
        ('NLP in Finance', 'This project demonstrates the power of modern NLP techniques, particularly transformer models like FinBERT, for extracting quantifiable insights from unstructured financial text.'),
        ('Automated Analysis', 'Manual reading of lengthy SEC filings is time-consuming. This pipeline automates sentiment extraction, enabling rapid analysis of large volumes of corporate communications.'),
        ('Multi-Model Validation', 'Using three different sentiment models provides robust analysis. Agreement across models increases confidence in the findings.'),
    ]
    
    for point_title, point_desc in importance_points:
        p = doc.add_paragraph()
        p.add_run(f'• {point_title}: ').bold = True
        p.add_run(point_desc)
    
    # Key Findings Section
    doc.add_heading('Key Findings', level=1)
    
    doc.add_paragraph(
        'Our analysis of 20 text chunks from 10 companies across 2022-2023 revealed:'
    )
    
    # Add findings table
    findings_table = doc.add_table(rows=4, cols=2)
    findings_table.style = 'Table Grid'
    
    findings_data = [
        ('Average VADER Score', '0.852 (Highly Positive)'),
        ('Average TextBlob Polarity', '0.137 (Slightly Positive)'),
        ('Average FinBERT Score', '0.650 (Positive)'),
        ('Sentiment Distribution', '65% Positive, 35% Neutral, 0% Negative'),
    ]
    
    for i, (metric, value) in enumerate(findings_data):
        row = findings_table.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Company Rankings
    doc.add_heading('Company Rankings by CEO Sentiment', level=2)
    
    rankings_table = doc.add_table(rows=11, cols=4)
    rankings_table.style = 'Table Grid'
    
    headers = ['Rank', 'Company', 'CEO', 'FinBERT Score']
    header_row = rankings_table.rows[0]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
        header_row.cells[i].paragraphs[0].runs[0].bold = True
    
    rankings_data = [
        ('1', 'ExxonMobil', 'Darren Woods', '1.000'),
        ('2', 'Microsoft', 'Satya Nadella', '1.000'),
        ('3', 'Alphabet', 'Sundar Pichai', '1.000'),
        ('4', 'Apple', 'Tim Cook', '1.000'),
        ('5', 'Nvidia', 'Jensen Huang', '0.500'),
        ('6', 'Tesla', 'Elon Musk', '0.500'),
        ('7', 'JPMorgan', 'Jamie Dimon', '0.500'),
        ('8', 'Meta', 'Mark Zuckerberg', '0.500'),
        ('9', 'Berkshire Hathaway', 'Warren Buffett', '0.500'),
        ('10', 'Amazon', 'Andy Jassy', '0.000'),
    ]
    
    for i, (rank, company, ceo, score) in enumerate(rankings_data):
        row = rankings_table.rows[i + 1]
        row.cells[0].text = rank
        row.cells[1].text = company
        row.cells[2].text = ceo
        row.cells[3].text = score
    
    doc.add_paragraph()
    
    # Visualizations Section
    doc.add_heading('Visualizations', level=1)
    
    charts_dir = 'output/charts'
    
    visualizations = [
        ('sentiment_by_company.png', 'Sentiment by Company', 
         'This bar chart shows the average FinBERT sentiment score for each company. Higher scores indicate more positive CEO communications. ExxonMobil, Microsoft, Alphabet, and Apple lead with perfectly positive sentiment scores.'),
        ('sentiment_trend.png', 'Sentiment Trend Over Time',
         'This line chart tracks how CEO sentiment has evolved over the analysis period. It reveals patterns in communication tone across different market conditions.'),
        ('sentiment_heatmap.png', 'Model Comparison Heatmap',
         'This heatmap provides a side-by-side comparison of all three sentiment models (VADER, TextBlob, FinBERT) across all companies, revealing where models agree and diverge.'),
        ('wordcloud_overall.png', 'Word Clouds',
         'These word clouds visualize the most frequently used words in positive vs. negative sentiment text chunks, highlighting key themes in CEO communications.'),
        ('summary_table.png', 'Summary Metrics Table',
         'A comprehensive table showing all metrics for each company, including CEO name, dominant sentiment, and scores from all three models.'),
    ]
    
    for filename, title_text, description in visualizations:
        filepath = os.path.join(charts_dir, filename)
        
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(description)
        
        if os.path.exists(filepath):
            doc.add_picture(filepath, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f'[Image not found: {filename}]')
        
        doc.add_paragraph()
    
    # Methodology Section
    doc.add_heading('Methodology', level=1)
    
    doc.add_heading('Data Source', level=2)
    doc.add_paragraph(
        'CEO shareholder letters were extracted from SEC EDGAR 10-K annual filings. '
        'The SEC EDGAR database is the official repository of company filings required '
        'by US securities law. 10-K filings contain comprehensive annual reports including '
        'letters from the CEO to shareholders.'
    )
    
    doc.add_heading('Text Processing Pipeline', level=2)
    processing_steps = [
        'Downloaded 10-K filings for each target company (2022-2023)',
        'Extracted CEO letter sections using pattern matching',
        'Removed boilerplate legal disclaimers and forward-looking statements',
        'Chunked text into paragraph-level segments for analysis',
        'Filtered paragraphs containing company-building keywords',
        'Applied sentiment analysis models to each chunk',
        'Aggregated scores by company and generated visualizations',
    ]
    
    for step in processing_steps:
        doc.add_paragraph(f'• {step}')
    
    # Technical Details
    doc.add_heading('Technical Implementation', level=1)
    
    tech_details = [
        ('Programming Language', 'Python 3.8+'),
        ('Web Scraping', 'Requests, BeautifulSoup4, Selenium'),
        ('NLP Libraries', 'NLTK, TextBlob, Transformers (HuggingFace)'),
        ('ML Model', 'ProsusAI/FinBERT (440MB transformer model)'),
        ('Visualization', 'Matplotlib, Seaborn, WordCloud'),
        ('Data Processing', 'Pandas, NumPy'),
        ('Report Generation', 'Quarto (for HTML), python-docx (for Word)'),
    ]
    
    tech_table = doc.add_table(rows=len(tech_details), cols=2)
    tech_table.style = 'Table Grid'
    
    for i, (component, tools) in enumerate(tech_details):
        row = tech_table.rows[i]
        row.cells[0].text = component
        row.cells[1].text = tools
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This sentiment analysis reveals that CEOs of major US companies generally communicate '
        'with positive or neutral sentiment when discussing their companies\' journeys and strategies. '
        'The absence of negative sentiment across all analyzed text suggests that shareholder letters '
        'are carefully crafted to maintain investor confidence while still acknowledging challenges.'
    )
    doc.add_paragraph(
        'The multi-model approach (VADER, TextBlob, FinBERT) provides robust validation of sentiment '
        'findings. FinBERT, being specifically trained on financial text, offers the most reliable '
        'sentiment classification for SEC filings.'
    )
    doc.add_paragraph(
        'This pipeline can be extended to analyze historical trends, compare sectors, or monitor '
        'sentiment changes in real-time as new filings are released.'
    )
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Generated by CEO Sentiment Analysis Pipeline • March 2026')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    footer.runs[0].font.size = Pt(10)
    
    # Save document
    output_path = 'CEO_Sentiment_Analysis_Report.docx'
    doc.save(output_path)
    print(f'Word document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
